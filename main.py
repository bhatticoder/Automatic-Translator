from fastapi import FastAPI, Request, UploadFile, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import requests
import PyPDF2
import docx
import io
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
import arabic_reshaper
from bidi.algorithm import get_display
import textwrap
from dotenv import load_dotenv
from docx import Document
import json
import uuid
import time 

load_dotenv()

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

HISTORY_FILE = "history.json"
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
TRANSLATE_URL = "https://translation.googleapis.com/language/translate/v2"

# Load history from file or initialize empty list
def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r") as f:
            return json.load(f)
    return []

def save_history(history):
    with open(HISTORY_FILE, "w") as f:
        json.dump(history, f, indent=4)

translation_history = load_history()

# Auto delete old history entries (older than 48 hours)
def cleanup_expired_history():
    global translation_history
    cutoff_time = time.time() - (48 * 60 * 60)
    translation_history = [h for h in translation_history if h.get("timestamp", time.time()) > cutoff_time]
    save_history(translation_history)

@app.get("/languages")
async def get_supported_languages():
    url = "https://translation.googleapis.com/language/translate/v2/languages"
    params = {"key": GOOGLE_API_KEY, "target": "en"}
    response = requests.get(url, params=params)
    result = response.json()
    languages = result.get("data", {}).get("languages", [])
    return JSONResponse(content=languages)

@app.post("/translate")
async def translate(request: Request):
    data = await request.json()
    source_text = data.get("text")
    source_lang = data.get("src_lang")
    target_lang = data.get("tgt_lang")

    if not source_text or not target_lang:
        return JSONResponse(status_code=400, content={"error": "Missing required fields"})

    payload = {
        "q": source_text,
        "target": target_lang,
        "format": "text",
        "key": GOOGLE_API_KEY
    }

    if source_lang != "auto":
        payload["source"] = source_lang

    response = requests.get(TRANSLATE_URL, params=payload)

    if response.status_code == 200:
        translated_text = response.json()["data"]["translations"][0]["translatedText"]
        
        entry = {
            "id": str(uuid.uuid4()),
            "type": "text",
            "from": source_lang,
            "to": target_lang,
            "original": source_text,
            "translated": translated_text,
            "timestamp": time.time()  # Store time for auto-deletion
        }
        translation_history.append(entry)
        save_history(translation_history)

        return {"translation": translated_text}
    else:
        return JSONResponse(status_code=500, content={"error": "Translation failed", "details": response.text})

@app.post("/translate_file")
async def translate_file(file: UploadFile, fileTargetLang: str = Form(...)):
    text = ""
    filename = file.filename

    if filename.endswith(".pdf"):
        file.file.seek(0)
        reader = PyPDF2.PdfReader(file.file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    elif filename.endswith(".docx"):
        file.file.seek(0)
        file_bytes = await file.read()
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        return JSONResponse(status_code=400, content={"error": "Unsupported file format"})

    if not text.strip():
        return JSONResponse(status_code=400, content={"error": "No text found in file"})

    detect_url = "https://translation.googleapis.com/language/translate/v2/detect"
    detect_params = {"q": text, "key": GOOGLE_API_KEY}
    detect_response = requests.get(detect_url, params=detect_params)
    detect_result = detect_response.json()
    detected_lang = "auto"
    try:
        detected_lang = detect_result["data"]["detections"][0][0]["language"]
    except (KeyError, IndexError):
        detected_lang = "auto"

    params = {
        "q": text,
        "target": fileTargetLang,
        "format": "text",
        "key": GOOGLE_API_KEY,
        "source": detected_lang
    }
    response = requests.get(TRANSLATE_URL, params=params)
    result = response.json()
    try:
        translated_text = result["data"]["translations"][0]["translatedText"]
    except (KeyError, IndexError):
        translated_text = "Translation failed."

    entry = {
        "id": str(uuid.uuid4()),
        "type": "file",
        "filename": filename,
        "from": detected_lang,
        "to": fileTargetLang,
        "original": text,
        "translated": translated_text,
        "timestamp": time.time()  # Store time for auto-deletion
    }
    translation_history.append(entry)
    save_history(translation_history)

    return JSONResponse(content={
        "translation": translated_text,
        "detected_lang": detected_lang,
        "history_id": entry["id"]
    })

@app.post("/delete_history")
async def delete_history(request: Request):
    data = await request.json()
    ids_to_delete = data.get("ids", [])
    global translation_history
    translation_history = [h for h in translation_history if h["id"] not in ids_to_delete]
    save_history(translation_history)
    return {"status": "success"}

@app.get("/download_pdf")
def download_pdf(text: str):
    buffer = io.BytesIO()
    font_path = os.path.join("static", "amiri-regular.ttf")
    pdfmetrics.registerFont(TTFont("Amiri", font_path))
    p = canvas.Canvas(buffer)
    p.setFont("Amiri", 18)

    left_margin = 40
    top_margin = 800
    line_height = 24
    bottom_margin = 40
    max_chars_per_line = 60

    text_lines = text.split('\n')
    y = top_margin

    for line in text_lines:
        try:
            reshaped_text = arabic_reshaper.reshape(line)
            bidi_text = get_display(reshaped_text)
            display_line = bidi_text
        except Exception:
            display_line = line

        wrapped_lines = textwrap.wrap(display_line, max_chars_per_line)
        for wrap_line in wrapped_lines:
            p.drawString(left_margin, y, wrap_line)
            y -= line_height
            if y < bottom_margin:
                p.showPage()
                p.setFont("Amiri", 18)
                y = top_margin

    p.save()
    buffer.seek(0)
    return StreamingResponse(buffer, media_type="application/pdf",
                             headers={"Content-Disposition": "attachment; filename=translated.pdf"})

@app.get("/download_docx")
def download_docx(text: str):
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": "attachment; filename=translated.docx"})

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/history", response_class=HTMLResponse)
async def history(request: Request):
    cleanup_expired_history()
    history = load_history()
    return templates.TemplateResponse("history.html", {
        "request": request,
        "history": history
    })

@app.get("/history/{entry_id}", response_class=HTMLResponse)
async def history_detail(request: Request, entry_id: str):
    history = load_history()
    entry = next((item for item in history if item["id"] == entry_id), None)
    if not entry:
        raise HTTPException(status_code=404, detail="History entry not found")
    return templates.TemplateResponse("history_detail.html", {
        "request": request,
        "entry": entry
    })

# Modified load_history to auto-delete after 48 hours
def load_history():
    now = time.time()
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r") as f:
            history = json.load(f)
        # Keep only items newer than 48 hours
        history = [h for h in history if now - h.get("created_at", now) < 48*3600]
        save_history(history)
        return history
    return []
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
